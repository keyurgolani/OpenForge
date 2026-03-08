"""
DOCX processor for OpenForge Knowledge System.

Processes uploaded Word documents through:
1. Text extraction using python-docx
2. Metadata extraction
3. Thumbnail generation (placeholder)
4. Chunking and embedding
"""
import logging
from pathlib import Path
from uuid import UUID
from typing import Optional

from openforge.config import get_settings
from openforge.core.content_processors.base import ContentProcessor, ProcessorResult

logger = logging.getLogger("openforge.docx_processor")


class DocxProcessor(ContentProcessor):
    """Process Word documents for knowledge storage and retrieval."""

    name = "docx"
    supported_types = [
        "application/msword",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ]
    supported_extensions = [".doc", ".docx"]

    def __init__(self):
        self.settings = get_settings()

    async def process(
        self,
        file_path: str,
        workspace_id: UUID,
        knowledge_id: Optional[UUID] = None,
        **kwargs,
    ) -> ProcessorResult:
        """
        Process a Word document.

        Args:
            file_path: Path to the DOCX file
            workspace_id: UUID of the workspace
            knowledge_id: Optional UUID of the knowledge entry
            **kwargs: Additional options

        Returns:
            ProcessorResult with extracted text and metadata
        """
        result = ProcessorResult(success=False)

        doc_path = Path(file_path)
        if not doc_path.exists():
            result.error = f"Document file not found: {file_path}"
            logger.error(f"Document file not found: {file_path}")
            return result

        try:
            # Extract text using python-docx
            text_content = await self._extract_text(doc_path)
            result.content = text_content
            result.extracted_text = text_content
            result.metadata = {"word_count": len(text_content.split())}

            # Generate thumbnail (placeholder for Word docs)
            result.thumbnail_path = await self._generate_thumbnail(doc_path, workspace_id, knowledge_id)

            result.success = True

            # Embed if knowledge_id provided
            if knowledge_id and text_content.strip():
                try:
                    from openforge.core.knowledge_processor import knowledge_processor
                    await knowledge_processor.process_knowledge(
                        knowledge_id=knowledge_id,
                        workspace_id=workspace_id,
                        content=text_content,
                        knowledge_type="docx",
                        title=result.ai_title,
                        tags=[],
                    )
                    result.embedded = True
                except Exception as e:
                    logger.warning(f"Failed to embed DOCX content: {e}")

        except Exception as e:
            logger.exception(f"Error processing DOCX {knowledge_id}: {e}")
            result.error = str(e)

        return result

    async def _extract_text(self, doc_path: Path) -> str:
        """Extract text from Word document."""
        try:
            import docx
            doc = docx.Document(str(doc_path))
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Also extract table content
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)

            return "\n\n".join(paragraphs)
        except ImportError:
            logger.warning("python-docx not installed, cannot extract DOCX text")
            return ""
        except Exception as e:
            logger.error(f"Failed to extract DOCX text: {e}")
            return ""

    async def _generate_thumbnail(
        self, doc_path: Path, workspace_id: UUID, knowledge_id: Optional[UUID]
    ) -> Optional[str]:
        """Generate a placeholder thumbnail for Word documents."""
        if not knowledge_id:
            return None

        try:
            from PIL import Image, ImageDraw, ImageFont

            workspace_dir = Path(self.settings.workspace_root) / str(workspace_id)
            thumbnail_dir = workspace_dir / "thumbnails"
            thumbnail_dir.mkdir(parents=True, exist_ok=True)

            thumbnail_path = thumbnail_dir / f"{knowledge_id}.webp"

            # Create a simple placeholder image
            img = Image.new("RGB", (200, 260), color=(255, 255, 255))
            draw = ImageDraw.Draw(img)

            # Draw Word-like icon
            draw.rectangle([20, 20, 180, 240], fill=(41, 86, 155), outline=(30, 60, 120), width=2)
            draw.text((70, 110), "W", fill=(255, 255, 255))
            draw.text((40, 150), "DOCX", fill=(255, 255, 255))

            img.save(thumbnail_path, "WEBP", quality=85)
            return f"thumbnails/{knowledge_id}.webp"
        except Exception as e:
            logger.warning(f"Failed to generate DOCX thumbnail: {e}")
            return None
