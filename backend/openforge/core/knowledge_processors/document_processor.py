"""Document Processing Pipeline.

1. Extract text as paragraphs via python-docx (with heading levels)
2. Extract metadata (author, title, word count, paragraph count)
3. Generate thumbnail via LibreOffice
4. Chunk text → embed → store in Qdrant
"""
from __future__ import annotations

import logging
import os
from pathlib import Path
from typing import Optional
from uuid import UUID

logger = logging.getLogger("openforge.processors.document")


class DocumentProcessor:
    """Complete document (DOCX) knowledge processing pipeline."""

    async def process(
        self,
        knowledge_id: UUID,
        file_path: str,
        workspace_id: UUID,
        db_session=None,
    ) -> dict:
        """Run the full document processing pipeline. Returns metadata dict."""
        result = {
            "metadata": {},
            "text": "",
        }

        # ── Step 1: Extract text ──
        try:
            result["text"] = self._extract_text(file_path)
        except Exception as e:
            logger.warning("Document text extraction failed for %s: %s", knowledge_id, e)

        # ── Step 2: Extract metadata ──
        try:
            result["metadata"] = self._extract_metadata(file_path)
        except Exception as e:
            logger.warning("Document metadata extraction failed for %s: %s", knowledge_id, e)

        # ── Step 3: Thumbnail ──
        thumbnail_path = None
        try:
            from openforge.config import get_settings
            from openforge.core.knowledge_processors.thumbnail_utils import generate_office_thumbnail

            settings = get_settings()
            thumbnails_dir = os.path.join(settings.uploads_root, "knowledge-thumbnails")
            os.makedirs(thumbnails_dir, exist_ok=True)
            thumb_file = os.path.join(thumbnails_dir, f"{knowledge_id}.webp")
            if generate_office_thumbnail(file_path, thumb_file):
                thumbnail_path = thumb_file
        except Exception as e:
            logger.warning("Document thumbnail generation failed for %s: %s", knowledge_id, e)

        # ── Step 4: Embed text ──
        if result["text"] and len(result["text"].strip()) >= 20:
            try:
                await self._embed_text(knowledge_id, workspace_id, result["text"])
            except Exception as e:
                logger.warning("Document text embedding failed for %s: %s", knowledge_id, e)

        metadata = result["metadata"]
        return {
            "thumbnail_path": thumbnail_path,
            "file_metadata": {
                "author": metadata.get("author"),
                "doc_title": metadata.get("title"),
                "word_count": metadata.get("word_count"),
                "paragraph_count": metadata.get("paragraph_count"),
                "section_count": metadata.get("section_count"),
            },
            "content": result["text"],
            "ai_title": metadata.get("title") or Path(file_path).stem.replace("_", " ").replace("-", " ").title(),
        }

    def _extract_text(self, file_path: str) -> str:
        """Extract text from DOCX with heading structure preserved."""
        import docx

        doc = docx.Document(file_path)
        parts = []

        heading_map = {
            "Heading 1": "# ",
            "Heading 2": "## ",
            "Heading 3": "### ",
            "Heading 4": "#### ",
            "Heading 5": "##### ",
            "Heading 6": "###### ",
        }

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""
            prefix = heading_map.get(style_name, "")

            if prefix:
                parts.append(f"\n{prefix}{text}\n")
            elif style_name.startswith("List"):
                parts.append(f"- {text}")
            else:
                parts.append(text)

        # Also extract text from tables
        for table in doc.tables:
            table_rows = []
            for row_idx, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                table_rows.append("| " + " | ".join(cells) + " |")
                if row_idx == 0:
                    table_rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
            if table_rows:
                parts.append("\n" + "\n".join(table_rows) + "\n")

        full_text = "\n".join(parts)
        return full_text[:100000]  # Limit to 100k chars

    def _extract_metadata(self, file_path: str) -> dict:
        """Extract DOCX metadata."""
        import docx

        doc = docx.Document(file_path)
        core = doc.core_properties

        # Count words
        word_count = 0
        paragraph_count = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraph_count += 1
                word_count += len(text.split())

        return {
            "author": core.author or "",
            "title": core.title or "",
            "word_count": word_count,
            "paragraph_count": paragraph_count,
            "section_count": len(doc.sections),
        }

    async def _embed_text(
        self, knowledge_id: UUID, workspace_id: UUID, text: str
    ) -> None:
        """Embed extracted text into openforge_knowledge collection."""
        from openforge.core.knowledge_processor import knowledge_processor

        await knowledge_processor.process_knowledge(
            knowledge_id=knowledge_id,
            workspace_id=workspace_id,
            content=text,
            knowledge_type="document",
            title=None,
            tags=[],
        )


document_processor = DocumentProcessor()
